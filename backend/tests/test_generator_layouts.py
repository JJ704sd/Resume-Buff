"""
Round 3 J: 5 套排版模板 — generator.py layout dispatcher 行为测试

测点:
  - TestLayoutDispatch        5 套 layout 都能生成合法 docx + 默认 = classic
  - TestLayoutVisuals         视觉差异 — 颜色 / 斜体 / 底纹 / ■ marker / table / 字号
  - TestLayoutInvalid         非法 template 抛 ValueError
  - TestLayoutBackwardCompat  preview / generate 不传 template 也能工作
  - TestLayoutVisualsR3M1     R3-M.1 新增 3 套模板差异化 config
  - TestLayoutConfigSchema    R3-M.2: 8 套 LAYOUT_CONFIG 必含 5 个可读性参数
  - TestHeadingHierarchy      R3-M.2: H1/H2 字号 = body * ratio(直接调 helper 验证)
  - TestSectionSpacing        R3-M.2: H1 段前后距 + H2 = 一半
  - TestBulletSpacing         R3-M.2: bullet / shaded / skill 段后距 = item_spacing_pt
  - TestMetaSpacing           R3-M.2: meta 行段后距 = meta_spacing_pt

测"核心逻辑",不测:
  - 单纯 dict.get 取值
  - URL 字面量
  - mock 自指

测点设计原则:每加一个 case 必问"它测了什么核心价值?"
  - H1/H2 字号映射到 docx run.font.size → 验证参数实际生效(集成)
  - 段间距映射到 docx paragraph_format.space_before/after → 验证参数实际生效(集成)
  - 不测"h1_ratio >= h2_ratio"(已在 TestLayoutConfigSchema 锁 schema,避免重复)
"""
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.generator import (
    LAYOUT_CONFIG,
    _add_h1,
    _add_h2,
    _add_bullet,
    _add_meta_line,
    _add_skill_line,
    _add_shaded_highlight,
    generate_resume_docx,
    preview_resume,
    render_docx,
    build_sections,
)


# Round 3 M.2: 8 套模板共享的 5 个可读性参数(由 helper 改造消费)
READABILITY_KEYS = {
    "h1_size_ratio",       # H1 字号 = body * ratio
    "h2_size_ratio",       # H2 字号 = body * ratio
    "section_spacing_pt",  # H1 段前/后距(元组)
    "meta_spacing_pt",     # meta 行段后距
    "item_spacing_pt",     # bullet 段后距
}


# ----------------------------------------------------------------------
# 辅助:从 docx zip 内部读取并断言
# ----------------------------------------------------------------------
def _read_xml(path: Path, member: str = "word/document.xml") -> str:
    with zipfile.ZipFile(path) as z:
        return z.read(member).decode("utf-8")


def _docx_color_values(path: Path) -> set[str]:
    """提取 docx 内的所有 w:color 值(uppercase)。黑色 000000 也算在结果里。"""
    xml = _read_xml(path)
    return {v.upper() for v in re.findall(r'<w:color\s+w:val="([0-9A-Fa-f]{6})"', xml)}


def _docx_has_shd(path: Path) -> bool:
    """docx 是否包含 w:shd 元素(technical 模板的项目 highlight 底纹)"""
    return "w:shd" in _read_xml(path)


def _docx_has_italic(path: Path) -> bool:
    """docx 是否包含 w:i 元素(minimal 模板所有行都不应斜体)"""
    return bool(re.search(r"<w:i\s*/>|<w:i\s", _read_xml(path)))


def _docx_contains_text(path: Path, marker: str) -> bool:
    """docx 是否包含指定文本(technical 模板的 ■ skill marker)"""
    return marker in _read_xml(path)


def _docx_has_table(path: Path) -> bool:
    """docx 是否包含 w:tbl 元素(two_column 模板的双栏布局)"""
    return "<w:tbl" in _read_xml(path)


def _docx_normal_style_pt(path: Path) -> float | None:
    """读取 word/styles.xml 里 Normal 样式的 w:sz 字号(half-pt → pt)"""
    styles_xml = _read_xml(path, "word/styles.xml")
    m = re.search(r'<w:style[^>]*w:styleId="Normal"[^>]*>.*?</w:style>', styles_xml, re.DOTALL)
    if not m:
        return None
    sz = re.search(r'<w:sz[^/]*w:val="(\d+)"', m.group(0))
    if not sz:
        return None
    return int(sz.group(1)) / 2


# ----------------------------------------------------------------------
# TestLayoutDispatch: 5 套 layout 都能生成合法 docx + 默认 = classic
# ----------------------------------------------------------------------
class TestLayoutDispatch:
    @pytest.mark.parametrize("template", list(LAYOUT_CONFIG.keys()))
    def test_layout_generates_valid_docx(self, tmp_path: Path, template: str):
        out = generate_resume_docx(
            target_role="tech_metric",
            output_dir=tmp_path,
            template=template,
        )
        # 文件落地 + 体积合理
        assert out.exists(), f"{template}: 文件未生成"
        assert out.stat().st_size > 5_000, f"{template}: 文件过小 ({out.stat().st_size}B)"
        # 合法 zip + 关键 xml 存在
        assert zipfile.is_zipfile(out), f"{template}: 不是合法 zip"
        with zipfile.ZipFile(out) as z:
            assert "word/document.xml" in z.namelist(), f"{template}: 缺 document.xml"
            assert "word/styles.xml" in z.namelist(), f"{template}: 缺 styles.xml"

    def test_default_template_equals_classic(self, tmp_path: Path):
        """不传 template 与显式 template='classic' 行为一致"""
        out_default = generate_resume_docx(target_role="tech_metric", output_dir=tmp_path)
        out_classic = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="classic"
        )
        # 同色集合 + 同斜体
        assert _docx_color_values(out_default) == _docx_color_values(out_classic)
        assert _docx_has_italic(out_default) == _docx_has_italic(out_classic)
        # 同字号
        assert _docx_normal_style_pt(out_default) == _docx_normal_style_pt(out_classic)


# ----------------------------------------------------------------------
# TestLayoutVisuals: 视觉差异 — 颜色 / 斜体 / 底纹 / ■ marker / table / 字号
# ----------------------------------------------------------------------
class TestLayoutVisuals:
    def test_classic_uses_color(self, tmp_path: Path):
        """classic 至少要有 1 个非黑 RGBColor(深蓝/灰/中灰等)"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="classic"
        )
        colors = _docx_color_values(out)
        non_black = {c for c in colors if c != "000000"}
        assert non_black, f"classic 应有非黑颜色,实际 {colors}"

    def test_minimal_no_color(self, tmp_path: Path):
        """minimal 全文黑 — 不写任何 w:color 元素"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="minimal"
        )
        colors = _docx_color_values(out)
        assert colors == set(), f"minimal 应无颜色,实际 {colors}"

    def test_minimal_no_italic(self, tmp_path: Path):
        """minimal 所有行都不斜体 — meta 行(项目时间)不上 italic"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="minimal"
        )
        assert not _docx_has_italic(out), "minimal 不应有 italic 元素"

    def test_technical_has_shaded_highlights(self, tmp_path: Path):
        """technical 项目 highlights 用浅灰底纹 — w:shd 元素存在"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="technical"
        )
        assert _docx_has_shd(out), "technical 应有 w:shd 底纹元素"

    def test_technical_skill_marker(self, tmp_path: Path):
        """technical 技能行前缀 ■ 标识"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="technical"
        )
        assert _docx_contains_text(out, "■"), "technical 技能行应含 ■ 标识"

    def test_two_column_has_table(self, tmp_path: Path):
        """two_column 下方双栏 — 必有 w:tbl"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="two_column"
        )
        assert _docx_has_table(out), "two_column 应有 w:tbl 元素"
        # 反向断言:classic 应无 table
        out_classic = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="classic"
        )
        assert not _docx_has_table(out_classic), "classic 不应有 table(回归保护)"

    def test_single_column_font_size(self, tmp_path: Path):
        """single_column Normal style 字号 = 10pt(行距紧凑、适合 1 页纸)"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="single_column"
        )
        size_pt = _docx_normal_style_pt(out)
        assert size_pt == pytest.approx(10.0), f"single_column Normal 应为 10pt,实际 {size_pt}pt"


# ----------------------------------------------------------------------
# TestLayoutInvalid: 非法 template 抛 ValueError
# ----------------------------------------------------------------------
class TestLayoutInvalid:
    def test_invalid_template_raises(self, tmp_path: Path):
        """template='foo' 不在 LAYOUT_CONFIG → ValueError"""
        with pytest.raises(ValueError, match="不支持的模板"):
            generate_resume_docx(
                target_role="tech_metric", output_dir=tmp_path, template="foo"
            )


# ----------------------------------------------------------------------
# TestLayoutBackwardCompat: preview / generate 不传 template 也能工作
# ----------------------------------------------------------------------
class TestLayoutBackwardCompat:
    def test_preview_resume_default_template(self):
        """preview_resume 不传 template → 默认 classic"""
        data = preview_resume(target_role="tech_metric")
        # Round 3 J: preview 返回里带 template 字段
        assert data.get("template") == "classic"
        # sections 仍正常返回
        assert len(data["sections"]) >= 1
        assert data["target_role"] == "tech_metric"

    def test_generate_resume_docx_default_template(self, tmp_path: Path):
        """generate_resume_docx 不传 template → 默认走 classic 渲染"""
        out = generate_resume_docx(target_role="tech_metric", output_dir=tmp_path)
        assert out.exists()
        # classic 视觉特征:有 italic(项目时间 meta 行)
        assert _docx_has_italic(out), "默认 template 应走 classic(classic 有 italic meta)"
        assert _docx_normal_style_pt(out) == pytest.approx(10.5)


# ----------------------------------------------------------------------
# TestLayoutVisualsR3M1: Round 3 M.1 新增 3 套模板的差异化 config 断言
# (dispatch 覆盖由现有 TestLayoutDispatch.test_layout_generates_valid_docx 参数化
#  自动扩展 — 该 parametrize 遍历 LAYOUT_CONFIG.keys(),新增 3 项自动覆盖 3 个新模板)
# ----------------------------------------------------------------------
class TestLayoutVisualsR3M1:
    def test_academic_larger_font_size(self):
        """academic font_size_body = 11.0(读博 / 出国申请 字号偏大)"""
        assert LAYOUT_CONFIG["academic"]["font_size_body"] == pytest.approx(11.0), (
            f"academic font_size_body 应为 11.0, 实际 {LAYOUT_CONFIG['academic']['font_size_body']}"
        )

    def test_internet_smaller_font_size(self):
        """internet font_size_body = 10.0(字节阿里 style 单页紧凑)"""
        assert LAYOUT_CONFIG["internet"]["font_size_body"] == pytest.approx(10.0), (
            f"internet font_size_body 应为 10.0, 实际 {LAYOUT_CONFIG['internet']['font_size_body']}"
        )

    def test_internet_has_skill_marker(self):
        """internet skill_marker = '▸ '(互联网简历技术感前缀)"""
        assert LAYOUT_CONFIG["internet"]["skill_marker"] == "▸ ", (
            f"internet skill_marker 应为 '▸ ', 实际 {LAYOUT_CONFIG['internet']['skill_marker']!r}"
        )

    def test_bilingual_default_margins(self):
        """bilingual margins 接近 2.0cm 四边(双语布局相对宽松)"""
        margins = LAYOUT_CONFIG["bilingual"]["margins_cm"]
        assert margins == pytest.approx((2.0, 2.0, 2.0, 2.0)), (
            f"bilingual margins 应为 (2.0, 2.0, 2.0, 2.0), 实际 {margins}"
        )


# ----------------------------------------------------------------------
# TestLayoutConfigSchema: Round 3 M.2 — 8 套 LAYOUT_CONFIG 必含 5 个可读性参数
# (锁 schema,任一模板缺关键 → fail,防 R3-M.2 之后的 round 改 config 时不小心删字段)
# ----------------------------------------------------------------------
class TestLayoutConfigSchema:
    def test_all_layouts_have_readability_keys(self):
        """8 套模板 LAYOUT_CONFIG 必含 5 个可读性参数(任一缺 → fail)"""
        for template, cfg in LAYOUT_CONFIG.items():
            missing = READABILITY_KEYS - set(cfg.keys())
            assert not missing, (
                f"{template} 缺可读性参数 {missing},需要补齐后才能被 helper 消费"
            )

    def test_h1_ratio_above_h2_ratio(self):
        """H1 size ratio 永远 >= H2 size ratio(8 套模板都满足,层次保证)"""
        for template, cfg in LAYOUT_CONFIG.items():
            assert cfg["h1_size_ratio"] >= cfg["h2_size_ratio"], (
                f"{template}: h1_ratio ({cfg['h1_size_ratio']}) < h2_ratio ({cfg['h2_size_ratio']})"
            )

    def test_h2_ratio_above_one(self):
        """H2 ratio >= 1.0(不会让 H2 比 body 还小,失去层次)"""
        for template, cfg in LAYOUT_CONFIG.items():
            assert cfg["h2_size_ratio"] >= 1.0, (
                f"{template}: h2_ratio ({cfg['h2_size_ratio']}) < 1.0(应至少等于 body 字号)"
            )

    def test_section_spacing_pt_is_two_tuple(self):
        """section_spacing_pt 必须是 (before, after) 2 元组(供 _add_h1 同时消费前后距)"""
        for template, cfg in LAYOUT_CONFIG.items():
            sp = cfg["section_spacing_pt"]
            assert isinstance(sp, tuple) and len(sp) == 2, (
                f"{template}: section_spacing_pt 必须是 (before, after) 2 元组, 实际 {sp!r}"
            )
            assert sp[0] >= 0 and sp[1] >= 0, (
                f"{template}: section_spacing_pt 必须 >= 0, 实际 {sp}"
            )

    def test_meta_and_item_spacing_non_negative(self):
        """meta_spacing_pt / item_spacing_pt 必须 >= 0(避免段后距为负导致行重叠)"""
        for template, cfg in LAYOUT_CONFIG.items():
            assert cfg["meta_spacing_pt"] >= 0, (
                f"{template}: meta_spacing_pt ({cfg['meta_spacing_pt']}) < 0"
            )
            assert cfg["item_spacing_pt"] >= 0, (
                f"{template}: item_spacing_pt ({cfg['item_spacing_pt']}) < 0"
            )


# ----------------------------------------------------------------------
# Round 3 M.2 — 标题层级 / 段间距 / bullet 段间距 参数化生效测试
# (直接调 helper 函数,验证 LAYOUT_CONFIG 的可读性参数真的写到 docx 上 — 核心集成价值)
# ----------------------------------------------------------------------
class TestHeadingHierarchy:
    """_add_h1 / _add_h2 字号 = body * ratio,验证 4 套模板差异化
    注:OOXML 把 w:sz 存为 half-point 整数,所以 12.6pt 会被存为 25 half-point=12.5pt。
    断言用 abs=0.5 容忍这个根本存储精度,而不是断言"完全相等"。
    """

    def test_h1_size_uses_h1_size_ratio_classic(self):
        """classic: H1 = body(10.5) * 1.20 = 12.6pt(±0.5 half-point 存储精度)"""
        doc = Document()
        p = _add_h1(doc, "教育背景", RGBColor(0, 0, 0), LAYOUT_CONFIG["classic"])
        actual_pt = p.runs[0].font.size.pt
        expected_pt = LAYOUT_CONFIG["classic"]["font_size_body"] * LAYOUT_CONFIG["classic"]["h1_size_ratio"]
        assert actual_pt == pytest.approx(expected_pt, abs=0.5), (
            f"classic H1 应为 {expected_pt}pt±0.5,实际 {actual_pt}pt"
        )

    def test_h2_size_uses_h2_size_ratio_classic(self):
        """classic: H2 = body(10.5) * 1.05 = 11.025pt(±0.5)"""
        doc = Document()
        p = _add_h2(doc, "项目名 | 角色", LAYOUT_CONFIG["classic"])
        actual_pt = p.runs[0].font.size.pt
        expected_pt = LAYOUT_CONFIG["classic"]["font_size_body"] * LAYOUT_CONFIG["classic"]["h2_size_ratio"]
        assert actual_pt == pytest.approx(expected_pt, abs=0.5), (
            f"classic H2 应为 {expected_pt}pt±0.5,实际 {actual_pt}pt"
        )

    def test_academic_h1_uses_academic_ratio(self):
        """academic: H1 = body(11) * 1.15 = 12.65pt(差异化,跟 classic 的 12.6pt 不同)"""
        doc = Document()
        p = _add_h1(doc, "教育背景", RGBColor(0, 0, 0), LAYOUT_CONFIG["academic"])
        actual_pt = p.runs[0].font.size.pt
        expected_pt = 11.0 * 1.15  # 12.65
        assert actual_pt == pytest.approx(expected_pt, abs=0.5), (
            f"academic H1 应为 {expected_pt}pt±0.5(差异化),实际 {actual_pt}pt"
        )

    def test_internet_h2_equals_body_size(self):
        """internet: H2 = body(10) * 1.0 = 10pt(同 body 字号,紧凑风格)"""
        doc = Document()
        p = _add_h2(doc, "项目名 | 角色", LAYOUT_CONFIG["internet"])
        actual_pt = p.runs[0].font.size.pt
        expected_pt = 10.0  # internet h2_ratio=1.0
        assert actual_pt == pytest.approx(expected_pt, abs=0.5), (
            f"internet H2 应等于 body {expected_pt}pt±0.5(层次靠粗体区分),实际 {actual_pt}pt"
        )


class TestSectionSpacing:
    """_add_h1 space_before/after = section_spacing_pt,_add_h2 = 一半"""

    def test_h1_section_spacing_pt_applied(self):
        """classic: H1 space_before=Pt(8), space_after=Pt(4)(来自 section_spacing_pt)"""
        doc = Document()
        p = _add_h1(doc, "X", RGBColor(0, 0, 0), LAYOUT_CONFIG["classic"])
        before, after = LAYOUT_CONFIG["classic"]["section_spacing_pt"]
        assert p.paragraph_format.space_before == Pt(before), (
            f"classic H1 space_before 应为 Pt({before}),实际 {p.paragraph_format.space_before}"
        )
        assert p.paragraph_format.space_after == Pt(after), (
            f"classic H1 space_after 应为 Pt({after}),实际 {p.paragraph_format.space_after}"
        )

    def test_h2_section_spacing_is_half_of_h1(self):
        """classic: H2 space = H1 一半(经典层次:大标题宽,小标题紧)"""
        doc = Document()
        p = _add_h2(doc, "X", LAYOUT_CONFIG["classic"])
        before, after = LAYOUT_CONFIG["classic"]["section_spacing_pt"]
        assert p.paragraph_format.space_before == Pt(before / 2), (
            f"H2 space_before 应为 H1 一半 Pt({before/2}),实际 {p.paragraph_format.space_before}"
        )
        assert p.paragraph_format.space_after == Pt(after / 2), (
            f"H2 space_after 应为 H1 一半 Pt({after/2}),实际 {p.paragraph_format.space_after}"
        )


class TestBulletSpacing:
    """_add_bullet / _add_shaded_highlight / _add_skill_line 段后距 = item_spacing_pt"""

    def test_bullet_classic_item_spacing_zero(self):
        """classic: item_spacing_pt=0 → bullet space_after=Pt(0)(R3-M.1 视觉保持)"""
        doc = Document()
        p = _add_bullet(doc, "X", LAYOUT_CONFIG["classic"])
        assert p.paragraph_format.space_after == Pt(0), (
            f"classic bullet space_after 应为 Pt(0),实际 {p.paragraph_format.space_after}"
        )

    def test_bullet_academic_item_spacing_two_pt(self):
        """academic: item_spacing_pt=2 → bullet space_after=Pt(2)(学术模板略散)"""
        doc = Document()
        p = _add_bullet(doc, "X", LAYOUT_CONFIG["academic"])
        expected = LAYOUT_CONFIG["academic"]["item_spacing_pt"]
        assert p.paragraph_format.space_after == Pt(expected), (
            f"academic bullet space_after 应为 Pt({expected}),实际 {p.paragraph_format.space_after}"
        )


class TestMetaSpacing:
    """_add_meta_line space_after = meta_spacing_pt(academic 3,其他 2)"""

    def test_meta_academic_spacing_three_pt(self):
        """academic: meta_spacing_pt=3 → meta 行 space_after=Pt(3)(学术 meta 略宽)"""
        doc = Document()
        p = _add_meta_line(doc, "2024.06 - 2024.09", LAYOUT_CONFIG["academic"])
        expected = LAYOUT_CONFIG["academic"]["meta_spacing_pt"]
        assert p.paragraph_format.space_after == Pt(expected), (
            f"academic meta space_after 应为 Pt({expected}),实际 {p.paragraph_format.space_after}"
        )


# ----------------------------------------------------------------------
# Round 3 M.2 — academic 专属 renderer
# (核心差异: 项目 highlights 简化 — 无 H2 项目名 / 无 meta line / 无 summary)
# ----------------------------------------------------------------------
from core.generator import _LAYOUT_DISPATCH, _render_academic  # noqa: E402


def _docx_paragraph_texts(path: Path) -> list[str]:
    """提取 docx 所有段落的纯文本(按出现顺序)"""
    doc = Document(str(path))
    return [p.text for p in doc.paragraphs]


def _docx_paragraphs_with_h2_size(path: Path, layout_cfg: dict) -> list[str]:
    """返回所有字号 = body * h2_size_ratio 的段落文本(H2 段识别)
    (帮助定位 H2 项目名 — classic 才有,academic 简化后无)
    """
    doc = Document(str(path))
    h2_pt = layout_cfg["font_size_body"] * layout_cfg["h2_size_ratio"]
    matches = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size and abs(r.font.size.pt - h2_pt) <= 0.5:
                if p.text.strip():
                    matches.append(p.text)
                break
    return matches


class TestAcademicRenderer:
    """academic 模板项目 highlights 简化 + 教育保持前置"""

    def test_academic_dispatch_uses_render_academic(self):
        """_LAYOUT_DISPATCH['academic'] 必须是 _render_academic(MVP 阶段是 _render_classic)"""
        assert _LAYOUT_DISPATCH["academic"] is _render_academic, (
            f"academic 模板应走 _render_academic, 实际 {_LAYOUT_DISPATCH['academic']!r}"
        )

    def test_academic_no_project_name_h2(self, tmp_path: Path):
        """academic 项目名不应作为 H2 段出现(classic 才有 '项目名 | 角色' H2)"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="academic"
        )
        h2_paragraphs = _docx_paragraphs_with_h2_size(out, LAYOUT_CONFIG["academic"])
        # tech_metric 4 个项目,classic 模板会出 4 个 "项目名 | 角色" H2
        # academic 简化后这 4 个 H2 都不应有
        for marker in ["某头部互联网公司医疗垂类大模型评测项目", "示例高校 × 某医疗器械公司", "DATAWHALE 开源社区", "大型体育赛事马拉松"]:
            for p_text in h2_paragraphs:
                assert marker not in p_text, (
                    f"academic 不应有项目名 H2,但发现 '{marker}' 在 H2 段 '{p_text}' 中"
                )

    def test_academic_no_project_period_meta(self, tmp_path: Path):
        """academic 项目时间(如 '2025.10 - 2025.12')不应作为独立 meta 段出现
        注:教育段时间嵌入 H2 段,honors 时间嵌在 'name (date)' 段里 — 都不算独立 meta 段;
        只有 classic 模板项目 group 里才把 period 作为独立段(整段就是 period 文本)。
        所以检查"段落文本完全等于 project period"即可锁定 _add_meta_line 是否被调用。
        """
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="academic"
        )
        all_texts = _docx_paragraph_texts(out)
        project_periods = ["2025.10 - 2025.12", "2026.1 - 至今", "2024.9 - 至今", "2024.12 - 2025.11"]
        for period in project_periods:
            for text in all_texts:
                assert text.strip() != period, (
                    f"academic 不应有项目时间独立 meta 段,但发现段落文本 = '{period}'"
                )

    def test_academic_no_project_summary(self, tmp_path: Path):
        """academic 项目概述(summary)不应出现(直接进 highlights)"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="academic"
        )
        all_texts = " | ".join(_docx_paragraph_texts(out))
        # 第一个项目的 summary 前缀(在 materials.json 里)
        summary_marker = "针对医疗垂直领域大语言模型,构建专业评测体系"
        assert summary_marker not in all_texts, (
            f"academic 不应有项目概述, 但发现 summary 内容"
        )

    def test_academic_education_before_projects(self, tmp_path: Path):
        """academic 模板中教育背景(H1 '教育背景')位置早于项目经历(H1 '项目经历')"""
        out = generate_resume_docx(
            target_role="tech_metric", output_dir=tmp_path, template="academic"
        )
        all_texts = _docx_paragraph_texts(out)
        edu_idx = next((i for i, t in enumerate(all_texts) if "教育背景" in t), -1)
        proj_idx = next((i for i, t in enumerate(all_texts) if "项目经历" in t), -1)
        assert edu_idx >= 0 and proj_idx >= 0, "教育背景 / 项目经历 段未找到"
        assert edu_idx < proj_idx, (
            f"academic 教育背景(idx={edu_idx})应在项目经历(idx={proj_idx})之前"
        )


# ----------------------------------------------------------------------
# Round 3 M.2 — 可读性扫描:8 套模板字号 / 段间距 / 行距 合理性
# (防止后续 round 改 LAYOUT_CONFIG 时引入过小字号 / 负段距 / 极端行距)
# ----------------------------------------------------------------------
class TestReadabilityAcrossLayouts:
    """8 套模板扫一遍 — 核心可读性 invariant,任一不满足 → fail"""

    def test_body_font_size_at_least_9pt(self):
        """8 套模板 body 字号 >= 9pt(可读性底线,再小就读不清)"""
        for template, cfg in LAYOUT_CONFIG.items():
            assert cfg["font_size_body"] >= 9.0, (
                f"{template}: font_size_body ({cfg['font_size_body']}) < 9pt — 不可读"
            )

    def test_meta_line_font_at_least_8pt(self):
        """8 套模板 meta 行(body * 0.88)>= 8pt(可读性底线)
        meta = font_size_body * 0.88(由 _add_meta_line 内部公式决定)
        """
        for template, cfg in LAYOUT_CONFIG.items():
            meta_pt = cfg["font_size_body"] * 0.88
            assert meta_pt >= 8.0, (
                f"{template}: meta ({meta_pt:.2f}pt) < 8pt — 不可读"
            )

    def test_all_spacings_non_negative(self):
        """8 套模板 4 个可读性间距参数全部 >= 0(段后距为负会导致行重叠)"""
        for template, cfg in LAYOUT_CONFIG.items():
            before, after = cfg["section_spacing_pt"]
            assert before >= 0 and after >= 0, (
                f"{template}: section_spacing_pt = {cfg['section_spacing_pt']} 含负值"
            )
            assert cfg["meta_spacing_pt"] >= 0, (
                f"{template}: meta_spacing_pt ({cfg['meta_spacing_pt']}) < 0"
            )
            assert cfg["item_spacing_pt"] >= 0, (
                f"{template}: item_spacing_pt ({cfg['item_spacing_pt']}) < 0"
            )

    def test_line_spacing_in_reasonable_range(self):
        """8 套模板 line_spacing 在 [1.15, 1.5](过紧/过松都不可读)"""
        for template, cfg in LAYOUT_CONFIG.items():
            ls = cfg["line_spacing"]
            assert 1.15 <= ls <= 1.5, (
                f"{template}: line_spacing ({ls}) 超出 [1.15, 1.5] 合理范围"
            )


# ----------------------------------------------------------------------
# Round 3 M.3 — academic_layout compact / detailed 双分支
# (验证 LAYOUT_CONFIG['academic']['academic_layout'] 字段被 _render_academic 真消费;
#  detailed 模式恢复 H2 项目名 + period meta + summary,compact 模式维持 R3-M.2 行为)
# ----------------------------------------------------------------------
def _generate_docx_with_layout(tmp_path: Path, layout_cfg: dict):
    """helper:用注入的 layout_cfg 走 _LAYOUT_DISPATCH['academic'] 生成 docx
    (绕开 generate_resume_docx 直接读 LAYOUT_CONFIG 的限制,允许测试注入 academic_layout)
    """
    from core.generator import _LAYOUT_DISPATCH, _setup_doc, ROLE_CONFIG
    out_dir = tmp_path / "out"
    out_dir.mkdir(exist_ok=True)
    sections = build_sections(target_role="tech_metric")
    doc = _setup_doc(layout_cfg)
    _LAYOUT_DISPATCH["academic"](doc, sections, ROLE_CONFIG["tech_metric"], layout_cfg)
    out = out_dir / "test.docx"
    doc.save(str(out))
    return out


class TestAcademicLayout:
    """academic_layout 字段驱动 _render_academic 项目段行为分支"""

    def test_default_is_compact(self):
        """LAYOUT_CONFIG['academic'] 默认 academic_layout = 'compact'(R3-M.2 行为兼容)"""
        assert LAYOUT_CONFIG["academic"].get("academic_layout") == "compact", (
            f"academic 默认 academic_layout 应为 'compact', 实际 {LAYOUT_CONFIG['academic'].get('academic_layout')!r}"
        )

    def test_compact_does_not_render_h2(self, tmp_path: Path):
        """academic_layout=compact:项目段无 H2 '项目名 | role'(compact 走 _render_project_group_academic_to 不渲染 H2)"""
        cfg = {**LAYOUT_CONFIG["academic"], "academic_layout": "compact"}
        out = _generate_docx_with_layout(tmp_path, cfg)
        h2_paragraphs = _docx_paragraphs_with_h2_size(out, cfg)
        for marker in ["某头部互联网公司医疗垂类大模型评测项目", "示例高校 × 某医疗器械公司"]:
            for p_text in h2_paragraphs:
                assert marker not in p_text, (
                    f"compact 不应有项目名 H2,但发现 '{marker}' 在 H2 段 '{p_text}' 中"
                )

    def test_compact_does_not_render_meta(self, tmp_path: Path):
        """academic_layout=compact:项目段无独立 period meta 段(整段文本 = period 视为 meta 段)"""
        cfg = {**LAYOUT_CONFIG["academic"], "academic_layout": "compact"}
        out = _generate_docx_with_layout(tmp_path, cfg)
        all_texts = _docx_paragraph_texts(out)
        periods = ["2025.10 - 2025.12", "2026.1 - 至今", "2024.9 - 至今", "2024.12 - 2025.11"]
        for period in periods:
            for text in all_texts:
                assert text.strip() != period, (
                    f"compact 不应有项目时间独立段,但发现段落 = '{period}'"
                )

    def test_detailed_renders_h2_project_name(self, tmp_path: Path):
        """academic_layout=detailed:项目段 H2 '项目名 | role' 必须出现(走 _render_project_group_academic_detailed_to)"""
        cfg = {**LAYOUT_CONFIG["academic"], "academic_layout": "detailed"}
        out = _generate_docx_with_layout(tmp_path, cfg)
        h2_paragraphs = _docx_paragraphs_with_h2_size(out, cfg)
        marker = "某头部互联网公司医疗垂类大模型评测项目"
        matched = [p_text for p_text in h2_paragraphs if marker in p_text]
        assert matched, (
            f"detailed 应有项目名 H2 含 '{marker}',但未在 H2 段中找到 (H2 段: {h2_paragraphs})"
        )

    def test_detailed_renders_period_meta(self, tmp_path: Path):
        """academic_layout=detailed:项目段 period 作为独立 meta 段(整段文本 = period)"""
        cfg = {**LAYOUT_CONFIG["academic"], "academic_layout": "detailed"}
        out = _generate_docx_with_layout(tmp_path, cfg)
        all_texts = _docx_paragraph_texts(out)
        periods = ["2025.10 - 2025.12", "2026.1 - 至今", "2024.9 - 至今", "2024.12 - 2025.11"]
        matched = [p for p in all_texts if p.strip() in periods]
        assert len(matched) >= 1, (
            f"detailed 应至少 1 个项目时间独立段,实际段落 = {all_texts}"
        )

    def test_detailed_renders_summary(self, tmp_path: Path):
        """academic_layout=detailed:项目段 summary 文本必须出现(走 _add_text)"""
        cfg = {**LAYOUT_CONFIG["academic"], "academic_layout": "detailed"}
        out = _generate_docx_with_layout(tmp_path, cfg)
        joined = " | ".join(_docx_paragraph_texts(out))
        summary_marker = "针对医疗垂直领域大语言模型,构建专业评测体系"
        assert summary_marker in joined, (
            f"detailed 应渲染 summary, 但未找到 '{summary_marker}'"
        )


# ----------------------------------------------------------------------
# Round 3 M.3 — bilingual 模板激活 + 3 个双语 helper + graceful degradation
# (验证 _render_bilingual_header/education/project_group_to 真消费 *_en 字段,
#  缺失时 graceful 降级到单语言 — 不抛异常,docx 仍合法)
# ----------------------------------------------------------------------
from core.generator import (
    _render_bilingual_header_to,
    _render_bilingual_education_to,
    _render_bilingual_project_group_to,
    _render_bilingual,
    _render_classic,
)  # noqa: E402


def _bilingual_mock_sections(include_en: bool = True) -> list:
    """构造 mock sections 喂给 3 个 bilingual helper(不需要 build_sections,测 helper 本体)
    include_en=True → name_en / school_en / title_en 都填,验证完整双语渲染
    include_en=False → _en 字段空字符串,验证 graceful 降级
    """
    from core.generator import Section
    en_suffix = "_en"
    name_en = "Mock Name EN" if include_en else ""
    school_en = "Mock University" if include_en else ""
    major_en = "Mock Major" if include_en else ""
    title_en = "Mock Project EN" if include_en else ""

    header = Section(
        type="header",
        title="个人信息",
        content={
            "name": "测试同学",
            "name_en": name_en,
            "intention": "AI 测试实习生",
            "contact": "13800000000 | test@example.com | 现居 深圳",
        },
    )
    edu = Section(
        type="education",
        title="教育背景",
        content={
            "line": "测试大学 · 测试学院 · 测试专业 · 本科  |  2024.9 - 2028.6(大二)",
            "school_en": school_en,
            "major_en": major_en,
            "courses": "、".join(["程序设计基础(C++)", "算法与数据结构"]),
            "highlights": ["测试亮点 A", "测试亮点 B"],
        },
    )
    proj = Section(
        type="project",
        title="测试项目",
        content={
            "role": "测试工程师",
            "period": "2025.10 - 2025.12",
            "summary": "测试项目概述。",
            "highlights": ["测试 highlight 1", "测试 highlight 2"],
        },
    )
    project_group = Section(
        type="project_group",
        title="项目经历",
        content={"projects": [
            {
                "id": "test_proj",
                "title": "测试项目",
                "title_en": title_en,
                "content": proj.content,
            }
        ]},
    )
    return [header, edu, project_group]


class TestBilingualHeader:
    """bilingual header helper — 中文 22pt + 可选英文 14pt italic"""

    def _make_doc(self):
        return Document()

    def test_header_renders_name(self):
        """中文姓名必须出现(22pt bold)"""
        doc = self._make_doc()
        s = _bilingual_mock_sections(include_en=False)[0]
        _render_bilingual_header_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert any("测试同学" in t for t in texts), f"bilingual header 应渲染中文姓名,实际 {texts}"

    def test_header_renders_name_en_when_present(self):
        """name_en 非空时,英文姓名斜体行必须出现(10pt italic 行)"""
        doc = self._make_doc()
        s = _bilingual_mock_sections(include_en=True)[0]
        _render_bilingual_header_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert any("Mock Name EN" in t for t in texts), (
            f"bilingual header 有 name_en 时应渲染英文行,实际 {texts}"
        )
        # 英文行有 italic
        en_runs = [r for p in doc.paragraphs for r in p.runs if "Mock Name EN" in p.text]
        assert en_runs and en_runs[0].italic, (
            f"英文姓名行应 italic,实际 runs = {en_runs}"
        )

    def test_header_graceful_no_en_when_absent(self):
        """name_en 缺失时(空字符串)不抛异常,只渲染中文行"""
        doc = self._make_doc()
        s = _bilingual_mock_sections(include_en=False)[0]
        # 不抛异常 + 英文行不存在
        _render_bilingual_header_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert not any("Mock Name EN" in t for t in texts), (
            f"无 name_en 时不应有英文行,实际 {texts}"
        )
        assert any("测试同学" in t for t in texts), "中文姓名应保留"

    def test_header_layout_center_default(self):
        """bilingual 默认 header_align=center → 姓名段对齐 = CENTER"""
        doc = self._make_doc()
        s = _bilingual_mock_sections(include_en=False)[0]
        _render_bilingual_header_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        # 找到含"测试同学"的段
        name_p = next(p for p in doc.paragraphs if "测试同学" in p.text)
        assert name_p.alignment == WD_ALIGN_PARAGRAPH.CENTER, (
            f"bilingual 默认居中对齐,实际 {name_p.alignment}"
        )


class TestBilingualEducation:
    """bilingual education helper — 中文 H2 + 可选英文 10pt italic"""

    def test_education_renders_line(self):
        """中文教育行(line)必须出现"""
        doc = Document()
        s = _bilingual_mock_sections(include_en=False)[1]
        _render_bilingual_education_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert any("测试大学" in t for t in texts), (
            f"bilingual education 应渲染中文教育行,实际 {texts}"
        )

    def test_education_renders_school_en_when_present(self):
        """school_en + major_en 非空时,英文斜体行存在(school_en | major_en)"""
        doc = Document()
        s = _bilingual_mock_sections(include_en=True)[1]
        _render_bilingual_education_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        en_line = next((t for t in texts if "Mock University" in t and "Mock Major" in t), None)
        assert en_line, f"应有英文教育行含 school_en | major_en,实际 {texts}"
        # 验证 italic
        en_runs = [r for p in doc.paragraphs for r in p.runs if en_line in p.text]
        assert en_runs and en_runs[0].italic, "英文教育行应 italic"

    def test_education_graceful_no_en_when_absent(self):
        """school_en / major_en 都缺失时,只渲染中文,无英文段"""
        doc = Document()
        s = _bilingual_mock_sections(include_en=False)[1]
        _render_bilingual_education_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert not any("Mock University" in t for t in texts), (
            f"无 _en 时不应有英文教育行,实际 {texts}"
        )


class TestBilingualProject:
    """bilingual project_group helper — 中文项目名 H2 + 可选英文副标题"""

    def test_project_renders_title(self, tmp_path: Path):
        """中文项目名 H2 必须出现"""
        doc = Document()
        s = _bilingual_mock_sections(include_en=False)[2]
        _render_bilingual_project_group_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        # 临时 save 让 _docx_paragraphs_with_h2_size 可读
        out = tmp_path / "test.docx"
        doc.save(str(out))
        h2_paragraphs = _docx_paragraphs_with_h2_size(out, LAYOUT_CONFIG["bilingual"])
        assert any("测试项目" in t and "测试工程师" in t for t in h2_paragraphs), (
            f"bilingual 项目应有 '测试项目 | 测试工程师' H2,实际 H2 段: {h2_paragraphs}"
        )

    def test_project_renders_title_en_when_present(self):
        """title_en 非空时,英文副标题斜体行存在(在 H2 之后,meta 之前)"""
        doc = Document()
        s = _bilingual_mock_sections(include_en=True)[2]
        _render_bilingual_project_group_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert any("Mock Project EN" in t for t in texts), (
            f"bilingual 项目有 title_en 应渲染英文副标题,实际 {texts}"
        )
        # 验证 italic
        en_runs = [r for p in doc.paragraphs for r in p.runs if "Mock Project EN" in p.text]
        assert en_runs and en_runs[0].italic, "英文项目副标题应 italic"

    def test_project_graceful_no_en_when_absent(self):
        """title_en 缺失时,无英文副标题段(直接 H2 → meta)"""
        doc = Document()
        s = _bilingual_mock_sections(include_en=False)[2]
        _render_bilingual_project_group_to(doc, s, RGBColor(0, 0, 0), LAYOUT_CONFIG["bilingual"])
        texts = [p.text for p in doc.paragraphs]
        assert not any("Mock Project EN" in t for t in texts), (
            f"无 title_en 时不应有英文副标题,实际 {texts}"
        )


class TestBilingualDispatch:
    """_render_bilingual 入口 + _LAYOUT_DISPATCH 切换 + classic 回归"""

    def test_bilingual_layout_dispatches_to_bilingual_renderer(self):
        """_LAYOUT_DISPATCH['bilingual'] 必须切到 _render_bilingual(dead code 激活)"""
        assert _LAYOUT_DISPATCH["bilingual"] is _render_bilingual, (
            f"bilingual 模板应走 _render_bilingual, 实际 {_LAYOUT_DISPATCH['bilingual']!r}"
        )

    def test_classic_layout_unchanged(self):
        """classic 仍走 _render_classic(回归保护,本轮 R3-M.3 不动 classic 路径)"""
        assert _LAYOUT_DISPATCH["classic"] is _render_classic, (
            f"classic 仍应走 _render_classic, 实际 {_LAYOUT_DISPATCH['classic']!r}"
        )