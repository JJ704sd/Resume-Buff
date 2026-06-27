"""
外部简历读取 — 解析 .docx / .pdf / .txt 字节流为段落 list (R3-G 新增)

设计要点:
  - **纯内存**:不落盘,接收 bytes → 返回 list[dict]。避免隐私数据落本地
  - **同构段落**:docx / pdf / txt 都返回 [{idx, text, is_heading, page?}] 结构
  - **is_heading 判定**:复用 _looks_like_heading 规则(项目经历 / 教育背景 / 个人技能 等)
  - **不依赖 LLM / 不联网**:仅做格式解析

公开 API:
  - read_docx(data)              -> 读 .docx 字节流
  - read_pdf(data)               -> 读 .pdf  字节流(每页 page 编号)
  - read_txt(data)               -> 读 .txt  字节流
  - parse_resume_bytes(filename, data)
                                 -> 根据后缀分派,返回顶层 dict(给 /api/resume/parse-external)

错误类型:
  - UnsupportedFormatError  (HTTP 415): 后缀不在 .docx / .pdf / .txt
  - EmptyResumeError        (HTTP 422): 字节流为 0 字节
  - ParseError              (HTTP 422): 格式坏/无法解析(由调用方捕获包成 422)
"""
import io
import re
from pathlib import PurePosixPath

# ----------------------------------------------------------------------
# 错误类型(给 API 层 HTTPException 用)
# ----------------------------------------------------------------------
class UnsupportedFormatError(ValueError):
    """不支持的文件类型 → HTTP 415"""


class EmptyResumeError(ValueError):
    """空文件 → HTTP 422"""


# ----------------------------------------------------------------------
# 标题判定 (复用 scripts/mvp_rewrite_resume.py 的规则)
# 决策:把 _HEADING_PATTERNS + _looks_like_heading 也搬过来,保证两处一致。
# ----------------------------------------------------------------------
_HEADING_PATTERNS = [
    r"^.{0,15}(项目经历|项目经验|教育背景|教育经历|工作经历|实习经历|个人技能|专业技能|自我评价|个人简介|个人概况|荣誉奖项|获奖情况|获奖经历|校园经历|社会实践|竞赛经历|科研经历|工作业绩|工作概要).{0,5}$",
    r"^#\s+",                                  # markdown 风格
    r"^【.+】$",                                # 【xxx】
]


def _looks_like_heading(text: str) -> bool:
    """判定单段文本是否像简历的章节标题。

    设计:同 mvp_rewrite_resume.py 完全一致,保证两处解析结果稳定。
    """
    if not text:
        return False
    stripped = text.strip()
    if not stripped:
        return False
    for pat in _HEADING_PATTERNS:
        if re.match(pat, stripped):
            return True
    return False


# ----------------------------------------------------------------------
# 单格式解析
# ----------------------------------------------------------------------
def read_docx(data: bytes) -> list[dict]:
    """读 .docx 字节流,返回 [{idx, text, is_heading}] 段落 list。

    注:
      - 表格里的内容 python-docx 不会自动遍历,这里只取 Document.paragraphs
      - 纯空段落仍保留 (text="", is_heading=False),保证 idx 与原文一致
      - 不抽取 docx 样式/图片(MVP 朴素段落输出)
    """
    from docx import Document  # python-docx — 局部 import 避免启动开销

    if not data:
        raise EmptyResumeError("docx 字节流为空")

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # python-docx 抛 PackageNotFoundError / ValueError 等
        raise ValueError(f"docx 解析失败: {e}") from e

    paras: list[dict] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if not text.strip():
            paras.append({"idx": i, "text": "", "is_heading": False})
        else:
            paras.append({
                "idx": i,
                "text": text,
                "is_heading": _looks_like_heading(text),
            })
    return paras


def read_pdf(data: bytes) -> list[dict]:
    """读 .pdf 字节流,返回 [{idx, text, is_heading, page}] 段落 list。

    注:
      - 按页切分:同一页内连续行合并为一个 paragraph
      - page 从 1 开始 (符合人类阅读习惯)
      - 跨页空行不算段落,只在每页开头计数
      - 不抽取 PDF 样式/图片(MVP 朴素段落输出)
    """
    import fitz  # PyMuPDF — 局部 import

    if not data:
        raise EmptyResumeError("pdf 字节流为空")

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:
        raise ValueError(f"pdf 解析失败: {e}") from e

    paras: list[dict] = []
    idx = 0
    page_count = doc.page_count
    try:
        for pi in range(page_count):
            page = doc[pi]
            text = page.get_text("text") or ""
            for line in text.splitlines():
                stripped = line.strip()
                if not stripped:
                    # 空行 → 跳过(MVP 不保留),避免连续空段
                    continue
                paras.append({
                    "idx": idx,
                    "text": line,
                    "is_heading": _looks_like_heading(line),
                    "page": pi + 1,  # 1-based
                })
                idx += 1
    finally:
        doc.close()

    return paras


def read_txt(data: bytes) -> list[dict]:
    """读 .txt 字节流,返回 [{idx, text, is_heading}] 段落 list。

    注:
      - 编码:优先 utf-8,失败 fallback gbk (中文简历常见)
      - 空行跳过(避免连续空段)
      - 长文本无长度上限,但单文件 5MB 限制由 API 层兜底
    """
    if not data:
        raise EmptyResumeError("txt 字节流为空")

    # 编码探测:utf-8 优先,失败再试 gbk
    text: str
    for enc in ("utf-8", "utf-8-sig", "gbk"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        # 兜底:errors=replace(同 mvp_rewrite_resume.read_resume 的 txt 分支)
        text = data.decode("utf-8", errors="replace")

    paras: list[dict] = []
    idx = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        paras.append({
            "idx": idx,
            "text": line,
            "is_heading": _looks_like_heading(line),
        })
        idx += 1
    return paras


# ----------------------------------------------------------------------
# 顶层分派
# ----------------------------------------------------------------------
_SUPPORTED_SUFFIXES = (".docx", ".pdf", ".txt")


def _detect_suffix(filename: str) -> str:
    """从文件名取后缀,小写化。"""
    return PurePosixPath(filename or "").suffix.lower()


def parse_resume_bytes(filename: str, data: bytes) -> dict:
    """根据 filename 后缀分派到 read_docx / read_pdf / read_txt。

    Returns:
        dict:
          - filename:        原文件名(只取 basename,不含路径)
          - size_bytes:      int
          - paragraphs:      list[{idx, text, is_heading, page?}]
          - page_count:      int | None   (仅 pdf)
          - note:            str          "解析成功,共 N 段(M 标题)"

    Raises:
        UnsupportedFormatError: filename 后缀不在 _SUPPORTED_SUFFIXES
        EmptyResumeError:       data 为 0 字节
        ValueError:             解析失败(包成 HTTP 422)
    """
    if not data:
        raise EmptyResumeError("文件为空")

    suffix = _detect_suffix(filename)
    if suffix not in _SUPPORTED_SUFFIXES:
        raise UnsupportedFormatError(
            f"不支持的文件类型: {suffix or '(无后缀)'} "
            f"(仅支持 .docx / .pdf / .txt)"
        )

    if suffix == ".docx":
        paragraphs = read_docx(data)
        page_count = None
    elif suffix == ".pdf":
        paragraphs = read_pdf(data)
        # page_count = paragraphs 里出现过的最大 page 值(段落可能跨页分布不均)
        pages = {p.get("page") for p in paragraphs if p.get("page") is not None}
        page_count = max(pages) if pages else 0
    else:  # .txt
        paragraphs = read_txt(data)
        page_count = None

    # 兜底:解析出来 0 段(可能 docx 全是表格,pdf 是扫描版)→ 当作解析失败
    if not paragraphs:
        raise ValueError(
            "未从文件中提取到任何段落 — 可能是扫描版 PDF / "
            "docx 全是表格 / 空文件"
        )

    heading_count = sum(1 for p in paragraphs if p.get("is_heading"))
    note = f"解析成功,共 {len(paragraphs)} 段({heading_count} 标题)"

    # filename 只取 basename(防止前端传完整路径)
    from pathlib import Path
    safe_name = Path(filename).name

    return {
        "filename": safe_name,
        "size_bytes": len(data),
        "paragraphs": paragraphs,
        "page_count": page_count,
        "note": note,
    }