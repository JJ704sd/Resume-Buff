"""
简历生成器核心逻辑

设计:
  - 把"构造 sections(结构化内容)"和"写 docx(物理排版)"完全解耦
  - preview()  和 generate_docx() 共享 build_sections() 的输出,保证预览 = 下载
  - Round 1: 规则化选项目 + 模板化排版(不接 LLM)
  - Round 2: 接 LLM 智能改写项目描述(改写层接在 build_sections 之前)
  - Round 3 J: 5 套排版模板 — classic / single_column / two_column / minimal / technical
"""
import json
import re
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.llm_rewriter import rewrite_highlights, is_llm_enabled
from core.jd_ranker import rank_projects, rank_highlights, rank_skill_groups


MATERIALS_PATH = Path(__file__).parent.parent / "data" / "materials.json"

# ----------------------------------------------------------------------
# 范围声明 (Round 2): 启用 6 个 role,覆盖产品 / 算法 / 测试 / 标注 / 度量 / 通用
# ----------------------------------------------------------------------
ENABLED_ROLES = ["tech_metric", "product", "algorithm", "data_annot", "test_qa", "general"]

ROLE_CONFIG = {
    # ---- Round 1 原保留 ----
    "tech_metric": {
        "intention": "大模型技术度量实习生",
        "preferred_project_ids": ["company_medical_eval", "university_ecg", "datawhale", "volunteer"],
        "skill_keys": ["ai_ml", "evaluation_metrics", "tools", "programming_languages", "documentation", "medical"],
        "self_eval_key": "tech_metric",
        "title_color": RGBColor(0x1F, 0x4E, 0x79),  # 深蓝
        "highlights_fallback": [],  # 空 → 默认 fallback 到 ["general"]
    },
    # ---- Round 2 新增 ----
    "product": {
        "intention": "AI 产品经理实习生",
        # 公司医疗评测项目有 product 视角的 highlights(场景矩阵 / Golden Set / Badcase 反馈)
        "preferred_project_ids": ["company_medical_eval", "datawhale", "volunteer"],
        "skill_keys": ["ai_ml", "documentation", "data", "tools"],
        "self_eval_key": "general",  # 暂无 product 专属,降级 general
        "title_color": RGBColor(0xC0, 0x50, 0x4D),  # 红
        "highlights_fallback": [],
    },
    "algorithm": {
        "intention": "医疗 AI 算法实习生",
        # 大学 ECG 项目有 algorithm 专属 highlights(模型复现 / 架构对比 / 训练全流程)
        "preferred_project_ids": ["university_ecg", "company_medical_eval", "datawhale"],
        "skill_keys": ["programming_languages", "ai_ml", "tools", "medical", "evaluation_metrics"],
        "self_eval_key": "general",
        "title_color": RGBColor(0x2E, 0x75, 0xB6),  # 蓝
        "highlights_fallback": [],
    },
    "data_annot": {
        "intention": "大模型数据标注实习生",
        # 公司医疗评测项目有 data_annot 视角(标注准确率 / SOP / 错误样本)
        "preferred_project_ids": ["company_medical_eval", "datawhale"],
        "skill_keys": ["documentation", "data", "ai_ml", "medical"],
        "self_eval_key": "general",
        "title_color": RGBColor(0x70, 0x80, 0x30),  # 橄榄绿
        "highlights_fallback": [],
    },
    "test_qa": {
        "intention": "AI 测试 / QA 实习生",
        # test_qa 没专属 highlights,fallback 到 tech_metric(最接近的测试能力),再降级 general
        "preferred_project_ids": ["company_medical_eval", "university_ecg", "datawhale"],
        "skill_keys": ["ai_ml", "evaluation_metrics", "documentation", "tools"],
        "self_eval_key": "general",
        "title_color": RGBColor(0x70, 0x33, 0x99),  # 紫
        "highlights_fallback": ["tech_metric"],
    },
    "general": {
        "intention": "日常实习(通用方向)",
        # 全部项目都至少有 general 降级
        "preferred_project_ids": ["company_medical_eval", "university_ecg", "datawhale", "volunteer"],
        "skill_keys": ["programming_languages", "ai_ml", "tools", "documentation", "data"],
        "self_eval_key": "general",
        "title_color": RGBColor(0x60, 0x60, 0x60),  # 灰
        "highlights_fallback": [],
    },
}

SKILL_LABEL = {
    "programming_languages": "编程语言",
    "ai_ml": "AI / 算法",
    "tools": "工程与工具",
    "evaluation_metrics": "评测指标",
    "medical": "医学素养",
    "documentation": "文档与协作",
    "data": "数据处理",
}


# ----------------------------------------------------------------------
# Round 3 J: 5 套排版模板配置
# ----------------------------------------------------------------------
# 视觉差异驱动表 — 各 _render_* 通过 layout_cfg 决定表现
# - use_color=False        → 全文黑,无任何彩色 RGBColor
# - header_align="left"    → 姓名 / 求职意向 / 联系方式 改左对齐
# - font_size_body         → _apply_layout_globals 写入 Normal style + 各 helper 默认字号
# - line_spacing           → _add_text / _add_bullet 行距
# - margins_cm             → _apply_layout_globals 应用 (top, bottom, left, right)
# - shaded_highlights=True → 项目 highlights 用浅灰底纹(_add_shaded_highlight)
# - skill_marker           → 非空时 skills 段前缀(技术感 "■ ")
# - two_column=True        → 走 _render_two_column,下方 2 栏 table
LAYOUT_CONFIG = {
    "classic": {
        "name": "经典",
        "description": "居中 header + 彩色 H1,通用排版",
        "use_color": True,
        "header_align": "center",
        "font_size_body": 10.5,
        "line_spacing": 1.3,
        "margins_cm": (1.8, 1.8, 2.0, 2.0),
        "two_column": False,
        "shaded_highlights": False,
        "skill_marker": "",
        # ---- Round 3 M.2: 可读性参数(R3-M.1 视觉行为完全兼容) ----
        "h1_size_ratio": 1.20,          # H1 = body * 1.20 → 12.6pt
        "h2_size_ratio": 1.05,          # H2 = body * 1.05 → 11.025pt
        "section_spacing_pt": (8, 4),   # H1 段前/后距(pt)
        "meta_spacing_pt": 2,           # meta 行段后距(pt)
        "item_spacing_pt": 0,           # bullet 段后距(pt)
    },
    "single_column": {
        "name": "单栏紧凑",
        "description": "字号偏小、行距紧,适合 1 页纸",
        "use_color": True,
        "header_align": "center",
        "font_size_body": 10.0,
        "line_spacing": 1.15,
        "margins_cm": (1.5, 1.5, 1.8, 1.8),
        "two_column": False,
        "shaded_highlights": False,
        "skill_marker": "",
        "h1_size_ratio": 1.20,          # 12.0pt
        "h2_size_ratio": 1.05,          # 10.5pt
        "section_spacing_pt": (8, 4),
        "meta_spacing_pt": 2,
        "item_spacing_pt": 0,
    },
    "two_column": {
        "name": "双栏",
        "description": "上方 header,下方 2 栏(技能左 / 项目右)",
        "use_color": True,
        "header_align": "center",
        "font_size_body": 10.5,
        "line_spacing": 1.25,
        "margins_cm": (1.8, 1.8, 2.0, 2.0),
        "two_column": True,
        "shaded_highlights": False,
        "skill_marker": "",
        "h1_size_ratio": 1.20,
        "h2_size_ratio": 1.05,
        "section_spacing_pt": (8, 4),
        "meta_spacing_pt": 2,
        "item_spacing_pt": 0,
    },
    "minimal": {
        "name": "极简",
        "description": "全黑白、无斜体、无装饰",
        "use_color": False,
        "header_align": "center",
        "font_size_body": 10.5,
        "line_spacing": 1.3,
        "margins_cm": (1.8, 1.8, 2.0, 2.0),
        "two_column": False,
        "shaded_highlights": False,
        "skill_marker": "",
        "h1_size_ratio": 1.20,
        "h2_size_ratio": 1.05,
        "section_spacing_pt": (8, 4),
        "meta_spacing_pt": 2,
        "item_spacing_pt": 0,
    },
    "technical": {
        "name": "技术感",
        "description": "header 左对齐,skills ■ 前缀,项目高亮带底纹",
        "use_color": True,
        "header_align": "left",
        "font_size_body": 10.5,
        "line_spacing": 1.3,
        "margins_cm": (1.8, 1.8, 2.0, 2.0),
        "two_column": False,
        "shaded_highlights": True,
        "skill_marker": "■ ",
        "h1_size_ratio": 1.20,
        "h2_size_ratio": 1.05,
        "section_spacing_pt": (8, 4),
        "meta_spacing_pt": 2,
        "item_spacing_pt": 0,
    },
    # ---- Round 3 M.1 新增 3 套模板 (MVP: 复用 _render_classic, R3-M.2 academic 加专属 renderer) ----
    "academic": {
        "name": "学术 CV",
        "description": "适合读博 / 出国申请,字号 11pt 行距 1.5 边距 2.5cm,教育背景优先",
        "use_color": True,
        "header_align": "center",
        "font_size_body": 11.0,
        "line_spacing": 1.5,
        "margins_cm": (2.5, 2.5, 2.5, 2.5),
        "two_column": False,
        "shaded_highlights": False,
        "skill_marker": "",
        "academic_mode": True,  # R3-M.2 加专属 renderer(简化 highlights + 教育前置)
        # R3-M.3: academic_layout 决定项目段走 compact(简化版,默认)还是 detailed(同 classic 完整版)
        # compact: 无 H2 项目名 / 无 period meta / 无 summary — 适合履历表 / 紧凑学术 CV
        # detailed: 恢复 H2 + period meta + summary — 适合 Research Statement / 学术 CV 详细版
        "academic_layout": "compact",
        # 学术 CV: H1 比例略小(big body 11pt 不需要 1.2 倍),段间距更宽更舒展
        "h1_size_ratio": 1.15,          # 12.65pt
        "h2_size_ratio": 1.05,          # 11.55pt
        "section_spacing_pt": (10, 5),  # 学术段间距略大
        "meta_spacing_pt": 3,           # 学术 meta 间距略大
        "item_spacing_pt": 2,           # 学术 bullet 略散(避免视觉拥挤)
    },
    "internet": {
        "name": "互联网简洁",
        "description": "字节阿里 style,字号 10pt 行距 1.2 边距 1.5cm 单栏紧凑",
        "use_color": True,
        "header_align": "left",
        "font_size_body": 10.0,
        "line_spacing": 1.2,
        "margins_cm": (1.5, 1.5, 1.5, 1.5),
        "two_column": False,
        "shaded_highlights": False,
        "skill_marker": "▸ ",
        # 互联网: H2 比例 = 1.0(跟 body 同字号,紧凑风格),段间距偏小
        "h1_size_ratio": 1.20,          # 12.0pt
        "h2_size_ratio": 1.00,          # 10.0pt(同 body,层次靠粗体区分)
        "section_spacing_pt": (6, 3),   # 紧凑
        "meta_spacing_pt": 2,
        "item_spacing_pt": 0,
    },
    "bilingual": {
        "name": "中英双语",
        "description": "header / 教育 / 项目双语,适合外企或海外岗位",
        "use_color": True,
        "header_align": "center",
        "font_size_body": 10.5,
        "line_spacing": 1.3,
        "margins_cm": (2.0, 2.0, 2.0, 2.0),
        "two_column": False,
        "shaded_highlights": False,
        "skill_marker": "",
        # R3-M.2: bilingual_mode flag 保持 dead code(本轮只做可读性参数化,
        # 双语 header / 教育 / 项目副标题留 R3-M.3 激活)。注意:这里**有意保留** flag,
        # 让 R3-M.3 实施时不用再补 schema, 也不被 TestLayoutConfigSchema 误判。
        "bilingual_mode": True,
        # 双语: H1 略小(避免双语标题视觉过重),bullet 段间距略大(中英文不同 baseline 拉开)
        "h1_size_ratio": 1.18,          # 12.39pt
        "h2_size_ratio": 1.05,          # 11.025pt
        "section_spacing_pt": (8, 4),
        "meta_spacing_pt": 2,
        "item_spacing_pt": 2,
    },
}


# ----------------------------------------------------------------------
# Section 数据类(可序列化为 JSON 给前端预览)
# ----------------------------------------------------------------------
@dataclass
class Section:
    type: str  # "header" | "education" | "project_group" | "skills" | "honors" | "self_eval"
    title: str
    content: dict = field(default_factory=dict)


# ----------------------------------------------------------------------
# 工具
# ----------------------------------------------------------------------
def load_materials() -> dict:
    with open(MATERIALS_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def _pick_highlights(project: dict, target_role: str, fallback_chain: Optional[list[str]] = None) -> list[str]:
    """
    按 fallback 链顺序挑选 highlights:
      1. target_role 自己的
      2. fallback_chain 里指定的角色(按顺序)
      3. 兜底 "general"
      4. 空列表(项目没数据)
    """
    h = project.get("highlights", {}) or {}
    if h.get(target_role):
        return h[target_role]
    chain = fallback_chain or []
    for role in chain:
        if h.get(role):
            return h[role]
    return h.get("general", [])


# ----------------------------------------------------------------------
# Round 3 I: JD-driven 上下文工具
# ----------------------------------------------------------------------
def _truncate_kw(kw: str, limit: int = 20) -> str:
    """截断关键词字符串(LLM prompt 注入限长,避免 token 涨)。"""
    if not kw:
        return ""
    s = str(kw)
    return s if len(s) <= limit else s[:limit]


def _build_jd_focus(parsed_jd: Optional[dict]) -> Optional[dict]:
    """
    从 parsed JD 构造 LLM prompt 用的 jd_focus dict。
    字段:
      - matched:         parsed["raw_keywords"] 全部(去重)+ 截断
      - missing:         当 parsed 有 matched 但缺 missing 字段时,留空列表(实际 missing 由 match_score 给)
                         这里 fallback 用 raw_keywords 自身(命中数 > 0 部分)
      - tier_required:   parsed["tier_info"]["required"]
      - tier_preferred:  parsed["tier_info"]["preferred"]
    注:parsed_jd 来自 parse_jd(),没有 missing 字段;若调用方想要 missing,
    应当用 match_score() 的结果而不是 parse_jd() 的结果。本函数是"无 context 时的最佳努力"。
    """
    if not parsed_jd:
        return None
    raw = parsed_jd.get("raw_keywords") or []
    tier = parsed_jd.get("tier_info") or {}
    return {
        "matched": [_truncate_kw(k) for k in raw],
        "missing": [],  # parse_jd 不给 missing;若需要由调用方用 match_score
        "tier_required": [_truncate_kw(k) for k in (tier.get("required") or [])],
        "tier_preferred": [_truncate_kw(k) for k in (tier.get("preferred") or [])],
    }


# ----------------------------------------------------------------------
# Sections 构造 (preview 和 docx 共用这一份)
# ----------------------------------------------------------------------
def build_sections(
    target_role: str,
    intention: Optional[str] = None,
    custom_project_ids: Optional[list[str]] = None,
    *,
    jd_context: Optional[dict] = None,
) -> list[Section]:
    """
    构造完整的简历 sections 列表(可序列化为 JSON 预览,也可喂给 docx 渲染)。

    Round 3 I 新增:可选 jd_context(来自 jd_parser.parse_jd(jd_text))。
    传 jd_context 时,会:
      1. 重排 project 列表(按命中 JD 关键词数倒序,tie 时维持 preferred_ids 原顺序)
      2. 项目内重排 highlights(命中数倒序,维持原顺序在 tie)
      3. 重排 skill group(组内 item 命中数总和倒序)
    不传 jd_context / jd_context 为空 dict → 完全走原路径(向后兼容,字节级一致)。
    """
    if target_role not in ROLE_CONFIG:
        raise ValueError(f"不支持的岗位: {target_role},可选: {list(ROLE_CONFIG.keys())}")
    if target_role not in ENABLED_ROLES:
        raise ValueError(
            f"岗位 {target_role} 暂未启用,当前已启用: {ENABLED_ROLES}"
        )

    materials = load_materials()
    role_cfg = ROLE_CONFIG[target_role]
    final_intention = intention or role_cfg["intention"]
    preferred_ids = custom_project_ids or role_cfg["preferred_project_ids"]

    sections: list[Section] = []

    # ----- Header -----
    b = materials["basics"]
    sections.append(Section(
        type="header",
        title=b["name"],
        content={
            "name": b["name"],
            "intention": final_intention,
            "contact": f"{b['phone']}  |  {b['email']}  |  现居 {b['location']}",
        },
    ))

    # ----- Education -----
    edu = materials["education"]
    sections.append(Section(
        type="education",
        title="教育背景",
        content={
            "line": f"{edu['school']} · {edu['college']} · {edu['major']} · {edu['degree']}  |  {edu['period']}({edu['year']})",
            "courses": "、".join(edu.get("core_courses", [])),
            "highlights": edu.get("highlights", []),
        },
    ))

    # ----- Projects -----
    proj_sections = []
    proj_map = {p["id"]: p for p in materials["projects"]}

    # Round 3 I: jd_context 非空时,按命中数重排 project 列表
    if jd_context:
        ordered_projects = rank_projects(
            [proj_map[pid] for pid in preferred_ids if pid in proj_map],
            jd_context,
            preferred_order=preferred_ids,
        )
        ordered_pids = [p["id"] for p in ordered_projects]
    else:
        ordered_pids = [pid for pid in preferred_ids if pid in proj_map]

    # Round 3 I: 构造 jd_focus(供 LLM 改写用),只在 jd_context 存在时构造
    jd_focus = _build_jd_focus(jd_context) if jd_context else None

    for pid in ordered_pids:
        p = proj_map[pid]
        proj_highlights = _pick_highlights(
            p, target_role, role_cfg.get("highlights_fallback")
        )
        # Round 3 I: 项目内 highlight 排序(命中数倒序,维持原顺序在 tie)
        if jd_context and proj_highlights:
            proj_highlights = rank_highlights(proj_highlights, jd_context)
        # Round 2 #3: LLM 智能改写 (无 key / 失败 → 静默回退原文,不破现有 API)
        if proj_highlights and is_llm_enabled():
            try:
                proj_highlights = rewrite_highlights(
                    proj_highlights,
                    target_role=target_role,
                    jd_text=final_intention,
                    jd_focus=jd_focus,
                )
            except Exception:
                pass  # 静默降级 — 高层 build_sections 仍返回原文
        proj_sections.append(Section(
            type="project",
            title=p["name"],
            content={
                "role": p["role"],
                "period": p["period"],
                "summary": p.get("summary", ""),
                "highlights": proj_highlights,
                "tags": p.get("tags", []),
            },
        ))
    sections.append(Section(type="project_group", title="项目经历", content={"projects": [asdict(s) for s in proj_sections]}))

    # ----- Skills -----
    skills_content = []
    skills = materials["skills"]
    # Round 3 I: jd_context 非空时,按组内 item 命中数重排 skill_keys
    if jd_context:
        ordered_skill_keys = rank_skill_groups(role_cfg["skill_keys"], materials, jd_context)
    else:
        ordered_skill_keys = role_cfg["skill_keys"]
    for key in ordered_skill_keys:
        items = skills.get(key, [])
        if not items:
            continue
        skills_content.append({
            "label": SKILL_LABEL.get(key, key),
            "items": items,
        })
    sections.append(Section(type="skills", title="相关技能", content={"groups": skills_content}))

    # ----- Honors & Certs -----
    honors = []
    for h in materials.get("honors", []):
        date = h.get("date", "")
        honors.append(f"{h['name']}  ({date})" if date else h["name"])
    for c in materials.get("certs", []):
        date = c.get("date", "")
        issuer = c.get("issuer", "")
        tail = f"  ·  {issuer}  ({date})" if date else f"  ·  {issuer}"
        honors.append(f"{c['name']}{tail}")
    sections.append(Section(type="honors", title="荣誉与证书", content={"items": honors}))

    # ----- Self Eval -----
    self_eval_key = role_cfg["self_eval_key"]
    sentences = materials.get("self_eval_versions", {}).get(self_eval_key) \
        or materials.get("self_eval_versions", {}).get("general", [])
    sections.append(Section(type="self_eval", title="自我评价", content={"sentences": sentences}))

    return sections


# ======================================================================
# docx 渲染 (Round 3 J: 5 套 layout dispatcher)
# ======================================================================
# 设计要点:
# - 所有 helper (容器 = doc 或 cell) 都通过 `container.add_paragraph()` 添加段,
#   python-docx 的 Document 和 _Cell 都支持 add_paragraph,行为一致。
# - 视觉差异完全由 layout_cfg 驱动 — _render_classic / _render_single_column /
#   _render_minimal / _render_technical 共用 _dispatch_section,只是 layout_cfg 不同。
# - _render_two_column 是结构差异,走 table 布局。

def _set_chinese_font(run, font_name: str = "微软雅黑", size_pt: float = 10.5):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _apply_layout_globals(doc: Document, layout_cfg: dict) -> Document:
    """应用 margin / 行距 / Normal style 默认字号 / 中文字体"""
    top, bottom, left, right = layout_cfg["margins_cm"]
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(layout_cfg["font_size_body"])
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return doc


def _setup_doc(layout_cfg: Optional[dict] = None) -> Document:
    """创建 Document 并应用 layout 全局参数。layout_cfg=None → classic"""
    if layout_cfg is None:
        layout_cfg = LAYOUT_CONFIG["classic"]
    doc = Document()
    _apply_layout_globals(doc, layout_cfg)
    return doc


# ---- 段落 helper(签名:container, ...)— container 可以是 Document 或 _Cell ----

def _add_h1(container, text: str, color: RGBColor, layout_cfg: dict):
    """section 大标题 (e.g. 教育背景) — minimal 时不设 color(默认黑)
    Round 3 M.2: 字号 = body * h1_size_ratio,段前后距 = section_spacing_pt
    """
    p = container.add_paragraph()
    run = p.add_run(text)
    h1_size = layout_cfg["font_size_body"] * layout_cfg["h1_size_ratio"]
    _set_chinese_font(run, size_pt=h1_size)
    run.bold = True
    if layout_cfg.get("use_color", True) and color is not None:
        run.font.color.rgb = color
    before, after = layout_cfg["section_spacing_pt"]
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def _add_h2(container, text: str, layout_cfg: dict):
    """子标题 (e.g. 项目名 + 角色)
    Round 3 M.2: 字号 = body * h2_size_ratio,段前后距 = section_spacing_pt 的一半
    """
    p = container.add_paragraph()
    run = p.add_run(text)
    h2_size = layout_cfg["font_size_body"] * layout_cfg["h2_size_ratio"]
    _set_chinese_font(run, size_pt=h2_size)
    run.bold = True
    before, after = layout_cfg["section_spacing_pt"]
    # H2 段间距 = H1 的一半(经典层次:大标题宽,小标题紧)
    p.paragraph_format.space_before = Pt(before / 2)
    p.paragraph_format.space_after = Pt(after / 2)
    return p


def _add_meta_line(container, text: str, layout_cfg: dict):
    """项目时间 / meta 行 — minimal 时不 italic 不上灰
    Round 3 M.2: 字号 = body * 0.88(原硬编码 9pt),space_after = meta_spacing_pt
    """
    p = container.add_paragraph()
    run = p.add_run(text)
    meta_size = layout_cfg["font_size_body"] * 0.88
    _set_chinese_font(run, size_pt=meta_size)
    if layout_cfg.get("use_color", True):
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(layout_cfg["meta_spacing_pt"])
    return p


def _add_text(container, text: str, layout_cfg: dict, size_pt: Optional[float] = None):
    """正文段"""
    p = container.add_paragraph()
    run = p.add_run(text)
    if size_pt is None:
        size_pt = layout_cfg["font_size_body"]
    _set_chinese_font(run, size_pt=size_pt)
    p.paragraph_format.line_spacing = layout_cfg.get("line_spacing", 1.3)
    return p


def _add_bullet(container, text: str, layout_cfg: dict, size_pt: Optional[float] = None):
    """普通 bullet
    Round 3 M.2: 段后距 = item_spacing_pt
    """
    if size_pt is None:
        size_pt = layout_cfg["font_size_body"]
    p = container.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_chinese_font(run, size_pt=size_pt)
    p.paragraph_format.line_spacing = layout_cfg.get("line_spacing", 1.25)
    p.paragraph_format.space_after = Pt(layout_cfg["item_spacing_pt"])
    return p


def _add_shaded_highlight(container, text: str, layout_cfg: dict, size_pt: Optional[float] = None):
    """浅灰底纹 bullet(technical template 专用)— 用 w:shd 元素 fill=EEEEEE
    Round 3 M.2: 段后距 = item_spacing_pt
    """
    if size_pt is None:
        size_pt = layout_cfg["font_size_body"]
    p = container.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_chinese_font(run, size_pt=size_pt)
    p.paragraph_format.line_spacing = layout_cfg.get("line_spacing", 1.25)
    # Add light gray shading to paragraph
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEEEEE")
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(layout_cfg["item_spacing_pt"])
    return p


def _add_skill_line(container, label: str, items: list, layout_cfg: dict):
    """技能行 — technical 模板用 marker (e.g. '■ ') 加前缀
    Round 3 M.2: 段后距 = item_spacing_pt
    """
    marker = layout_cfg.get("skill_marker", "") or ""
    p = container.add_paragraph()
    text = f"{label}: " + "、".join(items)
    run = p.add_run(marker + text)
    _set_chinese_font(run, size_pt=layout_cfg["font_size_body"])
    p.paragraph_format.line_spacing = layout_cfg.get("line_spacing", 1.3)
    p.paragraph_format.space_after = Pt(layout_cfg["item_spacing_pt"])
    return p


# ---- section 渲染(到 container)----

def _render_header_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    """header(姓名/求职意向/联系方式)— technical 时左对齐"""
    c = s.content
    if layout_cfg.get("header_align") == "left":
        align = WD_ALIGN_PARAGRAPH.LEFT
    else:
        align = WD_ALIGN_PARAGRAPH.CENTER

    # 姓名 — 始终大字号 bold,不设 color(默认黑)
    p = container.add_paragraph()
    p.alignment = align
    run = p.add_run(c["name"])
    _set_chinese_font(run, size_pt=22)
    run.bold = True
    p.paragraph_format.space_after = Pt(4)

    # 求职意向 — minimal 时不上色
    p = container.add_paragraph()
    p.alignment = align
    run = p.add_run(f"求职意向:{c['intention']}")
    _set_chinese_font(run, size_pt=11)
    if layout_cfg.get("use_color", True):
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)

    # 联系方式 — minimal 时不上灰
    p = container.add_paragraph()
    p.alignment = align
    run = p.add_run(c["contact"])
    _set_chinese_font(run, size_pt=10)
    if layout_cfg.get("use_color", True):
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(8)


def _render_education_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    _add_h1(container, s.title, color, layout_cfg)
    _add_h2(container, s.content["line"], layout_cfg)
    if s.content.get("courses"):
        _add_text(container, f"核心课程:{s.content['courses']}", layout_cfg)
    for h in s.content.get("highlights", []):
        _add_bullet(container, h, layout_cfg)


def _render_project_group_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    _add_h1(container, s.title, color, layout_cfg)
    for proj in s.content["projects"]:
        c = proj["content"]
        _add_h2(container, f"{proj['title']}  |  {c['role']}", layout_cfg)
        _add_meta_line(container, c["period"], layout_cfg)
        if c.get("summary"):
            _add_text(container, c["summary"], layout_cfg)
        for h in c.get("highlights", []):
            if layout_cfg.get("shaded_highlights", False):
                _add_shaded_highlight(container, h, layout_cfg)
            else:
                _add_bullet(container, h, layout_cfg)


def _render_skills_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    _add_h1(container, s.title, color, layout_cfg)
    for g in s.content["groups"]:
        _add_skill_line(container, g["label"], g["items"], layout_cfg)


def _render_honors_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    _add_h1(container, s.title, color, layout_cfg)
    for item in s.content["items"]:
        _add_bullet(container, item, layout_cfg)


def _render_self_eval_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    _add_h1(container, s.title, color, layout_cfg)
    for sent in s.content["sentences"]:
        _add_bullet(container, sent, layout_cfg)


def _dispatch_section(container, s: Section, color: RGBColor, layout_cfg: dict):
    """section → renderer 路由"""
    if s.type == "header":
        _render_header_to(container, s, color, layout_cfg)
    elif s.type == "education":
        _render_education_to(container, s, color, layout_cfg)
    elif s.type == "project_group":
        _render_project_group_to(container, s, color, layout_cfg)
    elif s.type == "skills":
        _render_skills_to(container, s, color, layout_cfg)
    elif s.type == "honors":
        _render_honors_to(container, s, color, layout_cfg)
    elif s.type == "self_eval":
        _render_self_eval_to(container, s, color, layout_cfg)


# ---- 5 套 layout renderer ----

def _render_classic(doc: Document, sections: list[Section], role_cfg: dict, layout_cfg: dict):
    """classic / single_column / minimal / technical 共用 body — 差异由 layout_cfg 驱动"""
    color = role_cfg["title_color"]
    for s in sections:
        _dispatch_section(doc, s, color, layout_cfg)


# 4 个 layout 共享 _render_classic 的 body — 用 alias 表达语义,便于将来差异扩展
_render_single_column = _render_classic
_render_minimal = _render_classic
_render_technical = _render_classic


def _render_project_group_academic_to(
    container, s: Section, color: RGBColor, layout_cfg: dict
):
    """academic 简化版 project_group:无 H2 项目名 + 无 meta line + 无 summary,直接列 highlights
    (学术 CV 偏好:读者只关心做了什么;项目名 / 时间 / 概述等元信息由 education 段承担)
    """
    _add_h1(container, s.title, color, layout_cfg)
    for proj in s.content["projects"]:
        c = proj["content"]
        # 不渲染 H2 "项目名 | role" 行
        # 不渲染 meta line (时间 / 周期)
        # 不渲染 summary
        for h in c.get("highlights", []):
            _add_bullet(container, h, layout_cfg)


def _render_project_group_academic_detailed_to(
    container, s: Section, color: RGBColor, layout_cfg: dict
):
    """academic detailed 模式 project_group:恢复 H2 项目名 + period meta + summary
    (行为同 _render_project_group_to,但保留 academic 的 helper — 走同一份 layout_cfg,
    满足 _render_academic 按 academic_layout 分支 dispatch 的对称性)
    """
    _add_h1(container, s.title, color, layout_cfg)
    for proj in s.content["projects"]:
        c = proj["content"]
        _add_h2(container, f"{proj['title']}  |  {c['role']}", layout_cfg)
        _add_meta_line(container, c["period"], layout_cfg)
        if c.get("summary"):
            _add_text(container, c["summary"], layout_cfg)
        for h in c.get("highlights", []):
            if layout_cfg.get("shaded_highlights", False):
                _add_shaded_highlight(container, h, layout_cfg)
            else:
                _add_bullet(container, h, layout_cfg)


def _render_academic(doc: Document, sections: list[Section], role_cfg: dict, layout_cfg: dict):
    """academic 专属 renderer — 学术 CV 简化 highlights + 教育前置
    行为:
      - education / skills / honors / self_eval / header 走 _dispatch_section(同 classic)
      - project_group 走 academic_layout 分支:
          compact  → _render_project_group_academic_to(无 H2 / 无 meta / 无 summary,默认)
          detailed → _render_project_group_academic_detailed_to(恢复 H2 + meta + summary)
    注:build_sections 已经把 education 放第一位,renderer 不用动顺序。
    """
    color = role_cfg["title_color"]
    academic_layout = layout_cfg.get("academic_layout", "compact")
    for s in sections:
        if s.type == "project_group":
            if academic_layout == "detailed":
                _render_project_group_academic_detailed_to(doc, s, color, layout_cfg)
            else:
                _render_project_group_academic_to(doc, s, color, layout_cfg)
        else:
            _dispatch_section(doc, s, color, layout_cfg)


# ---- Round 3 M.3: bilingual 模板专属 helper (激活 bilingual_mode dead code) ----
# 设计要点:
#   - 中文为主,英文为辅(14pt 姓名 / 10pt 学校项目副标题),缺失时 graceful 降级到单语言
#   - header / education / project_group 走 bilingual 版 helper,其余段走 _dispatch_section
#   - bilingual 模板走 _LAYOUT_DISPATCH['bilingual'] → _render_bilingual
def _render_bilingual_header_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    """bilingual header:中文姓名(22pt bold)+ 可选英文姓名(14pt italic)+ 求职意向 + 联系方式"""
    c = s.content
    if layout_cfg.get("header_align") == "left":
        align = WD_ALIGN_PARAGRAPH.LEFT
    else:
        align = WD_ALIGN_PARAGRAPH.CENTER

    # 姓名(中)— 22pt bold,同 classic
    p = container.add_paragraph()
    p.alignment = align
    run = p.add_run(c["name"])
    _set_chinese_font(run, size_pt=22)
    run.bold = True
    p.paragraph_format.space_after = Pt(2)

    # 姓名(英)— graceful: 没有 name_en 就不渲染
    name_en = (c.get("name_en") or "").strip()
    if name_en:
        p = container.add_paragraph()
        p.alignment = align
        run = p.add_run(name_en)
        _set_chinese_font(run, size_pt=14)
        run.italic = True
        p.paragraph_format.space_after = Pt(4)

    # 求职意向 — bilingual 默认有色
    p = container.add_paragraph()
    p.alignment = align
    run = p.add_run(f"求职意向:{c['intention']}")
    _set_chinese_font(run, size_pt=11)
    if layout_cfg.get("use_color", True):
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)

    # 联系方式
    p = container.add_paragraph()
    p.alignment = align
    run = p.add_run(c["contact"])
    _set_chinese_font(run, size_pt=10)
    if layout_cfg.get("use_color", True):
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(8)


def _render_bilingual_education_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    """bilingual education:中文教育信息 + 可选英文小字(school_en / major_en 非空时)"""
    _add_h1(container, s.title, color, layout_cfg)
    c = s.content
    line = c.get("line", "")
    school_en = (c.get("school_en") or "").strip()
    major_en = (c.get("major_en") or "").strip()
    if school_en or major_en:
        # 中文行(H2)
        _add_h2(container, line, layout_cfg)
        # 英文小字(10pt italic)
        en_line = " | ".join(x for x in [school_en, major_en] if x)
        p = container.add_paragraph()
        run = p.add_run(en_line)
        _set_chinese_font(run, size_pt=10)
        run.italic = True
        p.paragraph_format.space_after = Pt(2)
    else:
        # graceful: 无英文就单语言(同 classic)
        _add_h2(container, line, layout_cfg)

    if c.get("courses"):
        _add_text(container, f"核心课程:{c['courses']}", layout_cfg)
    for h in c.get("highlights", []):
        _add_bullet(container, h, layout_cfg)


def _render_bilingual_project_group_to(container, s: Section, color: RGBColor, layout_cfg: dict):
    """bilingual project:中文项目名 + role(H2)+ 可选英文副标题(title_en 非空时)+ period meta + summary + highlights
    注:title_en 放在 project dict 顶层(与 id/title 同级),见 build_sections Step 3 R3-M.3
    """
    _add_h1(container, s.title, color, layout_cfg)
    for proj in s.content["projects"]:
        c = proj["content"]
        # 项目名(中)+ role(H2)
        _add_h2(container, f"{proj['title']}  |  {c['role']}", layout_cfg)
        # 英文副标题 — graceful: 没有 title_en 就不渲染
        title_en = (proj.get("title_en") or "").strip()
        if title_en:
            p = container.add_paragraph()
            run = p.add_run(title_en)
            _set_chinese_font(run, size_pt=10)
            run.italic = True
            p.paragraph_format.space_after = Pt(2)
        # period meta
        _add_meta_line(container, c["period"], layout_cfg)
        # summary
        if c.get("summary"):
            _add_text(container, c["summary"], layout_cfg)
        # highlights
        for h in c.get("highlights", []):
            if layout_cfg.get("shaded_highlights", False):
                _add_shaded_highlight(container, h, layout_cfg)
            else:
                _add_bullet(container, h, layout_cfg)


def _render_bilingual(doc: Document, sections: list[Section], role_cfg: dict, layout_cfg: dict):
    """bilingual renderer:header/education/project_group 走双语版 helper,其余段走 _dispatch_section
    graceful degradation:缺少 *_en 字段时 helper 内部 .get(..., '') 兜底空字符串,
    自然只渲染中文(单语言降级,docx 仍合法)。
    """
    color = role_cfg["title_color"]
    for s in sections:
        if s.type == "header":
            _render_bilingual_header_to(doc, s, color, layout_cfg)
        elif s.type == "education":
            _render_bilingual_education_to(doc, s, color, layout_cfg)
        elif s.type == "project_group":
            _render_bilingual_project_group_to(doc, s, color, layout_cfg)
        else:
            _dispatch_section(doc, s, color, layout_cfg)


def _add_two_column_table(doc: Document, sections: list[Section], role_cfg: dict, layout_cfg: dict):
    """双栏布局:左栏 education/skills/honors,右栏 project_group/self_eval"""
    color = role_cfg["title_color"]
    left_types = {"education", "skills", "honors"}
    right_types = {"project_group", "self_eval"}
    left_sections = [s for s in sections if s.type in left_types]
    right_sections = [s for s in sections if s.type in right_types]

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # 列宽:页面宽(A4=21cm) 减去左右 margin 后均分
    _, _, ml, mr = layout_cfg["margins_cm"]
    page_width_cm = 21.0 - ml - mr
    col_width = Cm(page_width_cm / 2)

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = col_width
    right_cell.width = col_width

    # 清掉 cell 自带的空 paragraph,避免顶端留白
    for cell in (left_cell, right_cell):
        if cell.paragraphs:
            p_el = cell.paragraphs[0]._p
            p_el.getparent().remove(p_el)

    for s in left_sections:
        _dispatch_section(left_cell, s, color, layout_cfg)
    for s in right_sections:
        _dispatch_section(right_cell, s, color, layout_cfg)


def _render_two_column(doc: Document, sections: list[Section], role_cfg: dict, layout_cfg: dict):
    """双栏 layout:header 全宽 + 下方 2 栏 table"""
    color = role_cfg["title_color"]
    # header 整行渲染
    for s in sections:
        if s.type == "header":
            _dispatch_section(doc, s, color, layout_cfg)
            break
    # 其余 sections 进 2 栏 table
    _add_two_column_table(doc, sections, role_cfg, layout_cfg)


# ---- dispatcher 入口 ----

_LAYOUT_DISPATCH = {
    "classic": _render_classic,
    "single_column": _render_single_column,
    "two_column": _render_two_column,
    "minimal": _render_minimal,
    "technical": _render_technical,
    # ---- Round 3 M.2: academic 加专属 renderer(简化 highlights) ----
    "academic": _render_academic,
    # ---- Round 3 M.1 MVP: internet 仍走 _render_classic ----
    "internet": _render_classic,
    # ---- Round 3 M.3: bilingual dead code 激活,改走 _render_bilingual ----
    "bilingual": _render_bilingual,
}


def render_docx(
    sections: list[Section],
    target_role: str,
    output_dir: Path,
    template: str = "classic",
) -> Path:
    """根据 sections 生成 .docx(template: classic / single_column / two_column / minimal / technical)"""
    if template not in LAYOUT_CONFIG:
        raise ValueError(f"不支持的模板: {template},可选: {list(LAYOUT_CONFIG.keys())}")
    if target_role not in ROLE_CONFIG:
        raise ValueError(f"不支持的岗位: {target_role},可选: {list(ROLE_CONFIG.keys())}")

    materials = load_materials()
    role_cfg = ROLE_CONFIG[target_role]
    layout_cfg = LAYOUT_CONFIG[template]

    intention = next((s.content["intention"] for s in sections if s.type == "header"), role_cfg["intention"])

    output_dir.mkdir(parents=True, exist_ok=True)
    doc = _setup_doc(layout_cfg)

    renderer = _LAYOUT_DISPATCH[template]
    renderer(doc, sections, role_cfg, layout_cfg)

    safe = re.sub(r"[^\w\u4e00-\u9fff]+", "_", intention)
    filename = f"{materials['basics']['name']}_{safe}_{template}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = output_dir / filename
    doc.save(str(out_path))
    return out_path


# ----------------------------------------------------------------------
# 公开 API
# ----------------------------------------------------------------------
def preview_resume(
    target_role: str,
    intention: Optional[str] = None,
    custom_project_ids: Optional[list[str]] = None,
    template: str = "classic",
    *,
    jd_text: Optional[str] = None,
) -> dict:
    """
    返回结构化预览(JSON 友好)。template 仅用于校验 / 透传到 docx 阶段(preview 不渲染 docx)。

    Round 3 I: jd_text 非空时,parse_jd 一次后传 jd_context 给 build_sections,
    触发 project / highlight / skill group 按命中数重排;
    返回额外字段 jd_match_counts(每 project / skill group 的命中关键词数),
    仅在 jd_text 非空时存在(无 jd 时键也存在但值为 None,前端按 None 不显示角标)。
    """
    if template not in LAYOUT_CONFIG:
        raise ValueError(f"不支持的模板: {template},可选: {list(LAYOUT_CONFIG.keys())}")

    jd_context = _resolve_jd_context(jd_text)
    sections = build_sections(
        target_role, intention, custom_project_ids, jd_context=jd_context
    )

    out: dict = {
        "target_role": target_role,
        "template": template,
        "intention": next((s.content["intention"] for s in sections if s.type == "header"), ""),
        "sections": [asdict(s) for s in sections],
        "jd_match_counts": _build_jd_match_counts(sections, jd_context) if jd_context else None,
    }
    return out


def generate_resume_docx(
    target_role: str,
    intention: Optional[str] = None,
    custom_project_ids: Optional[list[str]] = None,
    output_dir: Path = Path("output"),
    template: str = "classic",
    *,
    jd_text: Optional[str] = None,
) -> Path:
    """
    生成定制版简历 .docx(供 preview 确认后调用)。

    Round 3 I: jd_text 非空时透传到 build_sections,排序逻辑同上。
    """
    jd_context = _resolve_jd_context(jd_text)
    sections = build_sections(
        target_role, intention, custom_project_ids, jd_context=jd_context
    )
    return render_docx(sections, target_role, output_dir, template=template)


def _resolve_jd_context(jd_text: Optional[str]) -> Optional[dict]:
    """
    把 jd_text 原文解析成 parsed_jd dict 供 build_sections 使用。
    - jd_text 空 / None → None(走原路径,字节级一致)
    - 非空 → 调 parse_jd
    """
    if not jd_text or not jd_text.strip():
        return None
    # 局部 import 避免循环 + 单元测试时 jd_parser 已可用
    from core.jd_parser import parse_jd
    return parse_jd(jd_text)


def _build_jd_match_counts(
    sections: list[Section], jd_context: dict
) -> dict:
    """
    给 preview 返回追加 jd_match_counts:{projects:[N,N,...], skill_groups:[N,N,...]}
    数字 = 该 project / skill group 命中 JD 关键词的总数(供前端"命中 N 关键词"角标用)。
    """
    keywords: list[str] = jd_context.get("raw_keywords") or []

    projects_count: list[int] = []
    skill_count: list[int] = []

    for s in sections:
        if s.type == "project_group":
            for proj in s.content.get("projects", []):
                # 拼接项目内所有 highlight 文本(已按 rank_highlights 重排过)
                hl = proj.get("content", {}).get("highlights", []) or []
                joined = "\n".join(str(x) for x in hl)
                projects_count.append(_count_matches_inline(joined, keywords))
        elif s.type == "skills":
            for g in s.content.get("groups", []):
                joined = "\n".join(str(x) for x in g.get("items", []) or [])
                skill_count.append(_count_matches_inline(joined, keywords))

    return {"projects": projects_count, "skill_groups": skill_count}


def _count_matches_inline(text: str, keywords: list[str]) -> int:
    """跟 jd_ranker._count_matches 同语义(本地复制避免循环 import)。"""
    if not text or not keywords:
        return 0
    n = 0
    for kw in keywords:
        if kw and kw in text:
            n += 1
    return n